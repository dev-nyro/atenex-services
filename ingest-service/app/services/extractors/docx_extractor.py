import io
import docx  # python-docx
import structlog

log = structlog.get_logger(__name__)

class DocxExtractionError(Exception):
    pass

def extract_text_from_docx(file_bytes: bytes, filename: str = "unknown.docx") -> str:
    """Extrae texto de los bytes de un DOCX preservando saltos de párrafo."""
    log.debug("Extracting text from DOCX bytes", filename=filename)
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        text = "\n".join([p.text for p in doc.paragraphs])
        log.info("DOCX extraction successful", filename=filename, num_paragraphs=len(doc.paragraphs))
        return text
    except Exception as e:
        log.error("DOCX extraction failed", filename=filename, error=str(e))
        raise DocxExtractionError(f"Error extracting DOCX: {filename}") from e
