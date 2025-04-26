# ingest-service/app/services/ingest_pipeline.py
from __future__ import annotations

import os
import re
import pathlib
import uuid
import markdown # type: ignore
import html2text # type: ignore
import structlog
from typing import List, Dict, Any, Callable, Optional

# Direct library imports
from bs4 import BeautifulSoup
from pypdf import PdfReader
from docx import Document as DocxDocument
# LLM_FLAG: REMOVED - Global import and initialization of TextEmbedding removed from here.
from pymilvus import (
    Collection, CollectionSchema, FieldSchema, DataType, connections,
    utility, MilvusException
)
# LLM_FLAG: ADDED - Import only for type hinting within the function signature.
from fastembed.embedding import TextEmbedding

# Local application imports
from app.core.config import settings

log = structlog.get_logger(__name__)

# --------------- 1. TEXT EXTRACTION FUNCTIONS -----------------
# LLM_FLAG: NO_CHANGE - Text extraction functions remain the same.
def _extract_from_pdf(file_path: pathlib.Path) -> str:
    log.debug("Extracting text from PDF", path=str(file_path))
    try:
        reader = PdfReader(str(file_path))
        text_content = "\n\n".join(page.extract_text() or "" for page in reader.pages)
        log.info("PDF extraction successful", path=str(file_path), num_pages=len(reader.pages))
        return text_content
    except Exception as e:
        log.error("Failed to extract text from PDF", path=str(file_path), error=str(e), exc_info=True)
        raise ValueError(f"Error processing PDF {file_path.name}: {e}") from e

def _extract_from_docx(file_path: pathlib.Path) -> str:
    log.debug("Extracting text from DOCX", path=str(file_path))
    try:
        doc = DocxDocument(str(file_path))
        text_content = "\n\n".join(p.text for p in doc.paragraphs if p.text)
        log.info("DOCX extraction successful", path=str(file_path), num_paragraphs=len(doc.paragraphs))
        return text_content
    except Exception as e:
        log.error("Failed to extract text from DOCX", path=str(file_path), error=str(e), exc_info=True)
        raise ValueError(f"Error processing DOCX {file_path.name}: {e}") from e

def _extract_from_html(file_path: pathlib.Path) -> str:
    log.debug("Extracting text from HTML", path=str(file_path))
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f, "html.parser")
        for script_or_style in soup(["script", "style"]):
            script_or_style.decompose()
        text_content = soup.get_text(separator="\n", strip=True)
        log.info("HTML extraction successful", path=str(file_path))
        return text_content
    except Exception as e:
        log.error("Failed to extract text from HTML", path=str(file_path), error=str(e), exc_info=True)
        raise ValueError(f"Error processing HTML {file_path.name}: {e}") from e

def _extract_from_md(file_path: pathlib.Path) -> str:
    log.debug("Extracting text from Markdown", path=str(file_path))
    try:
        md_content = file_path.read_text(encoding="utf-8")
        html_content = markdown.markdown(md_content)
        h = html2text.HTML2Text()
        h.ignore_links = True
        h.ignore_images = True
        text_content = h.handle(html_content)
        log.info("Markdown extraction successful", path=str(file_path))
        return text_content
    except Exception as e:
        log.error("Failed to extract text from Markdown", path=str(file_path), error=str(e), exc_info=True)
        raise ValueError(f"Error processing Markdown {file_path.name}: {e}") from e

def _extract_from_txt(file_path: pathlib.Path) -> str:
    log.debug("Extracting text from TXT", path=str(file_path))
    try:
        text_content = file_path.read_text(encoding="utf-8")
        log.info("TXT extraction successful", path=str(file_path))
        return text_content
    except Exception as e:
        log.error("Failed to extract text from TXT", path=str(file_path), error=str(e), exc_info=True)
        raise ValueError(f"Error processing TXT {file_path.name}: {e}") from e

# LLM_FLAG: NO_CHANGE - Keep extractor map
EXTRACTORS: Dict[str, Callable[[pathlib.Path], str]] = {
    "application/pdf": _extract_from_pdf,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": _extract_from_docx,
    "application/msword": _extract_from_docx,
    "text/plain": _extract_from_txt,
    "text/markdown": _extract_from_md,
    "text/html": _extract_from_html,
}

# --------------- 2. CHUNKER FUNCTION -----------------
# LLM_FLAG: NO_CHANGE - Keep chunker function as is.
def chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> List[str]:
    if not text or text.isspace():
        return []
    words = re.split(r'\s+', text.strip())
    words = [word for word in words if word]

    if not words:
        return []

    chunks: List[str] = []
    current_pos = 0
    while current_pos < len(words):
        end_pos = min(current_pos + chunk_size, len(words))
        chunk = " ".join(words[current_pos:end_pos])
        chunks.append(chunk)

        current_pos += chunk_size - chunk_overlap
        if current_pos <= (end_pos - chunk_size):
             current_pos = end_pos - chunk_overlap
        if current_pos >= len(words):
            break
        current_pos = max(0, current_pos)

    log.debug(f"Chunked text into {len(chunks)} chunks.", requested_size=chunk_size, overlap=chunk_overlap)
    return chunks

# --------------- 3. MILVUS HELPERS (using pymilvus) -----------------
# LLM_FLAG: NO_CHANGE - Keep Milvus constants and helpers.
MILVUS_COLLECTION_NAME = settings.MILVUS_COLLECTION_NAME
MILVUS_EMBEDDING_DIM = settings.EMBEDDING_DIMENSION
MILVUS_PK_FIELD = "pk_id"
MILVUS_VECTOR_FIELD = "embedding"
MILVUS_CONTENT_FIELD = "content"
MILVUS_COMPANY_ID_FIELD = "company_id"
MILVUS_DOCUMENT_ID_FIELD = "document_id"
MILVUS_FILENAME_FIELD = "file_name"

_milvus_collection: Optional[Collection] = None
_milvus_connected = False

def _ensure_milvus_connection_and_collection() -> Collection:
    global _milvus_collection, _milvus_connected
    alias = "default"

    if not _milvus_connected:
        uri = settings.MILVUS_URI
        log.info("Connecting to Milvus...", uri=uri, alias=alias, timeout=settings.MILVUS_GRPC_TIMEOUT)
        try:
            connections.connect(alias=alias, uri=uri, timeout=settings.MILVUS_GRPC_TIMEOUT)
            _milvus_connected = True
            log.info("Successfully connected to Milvus.")
        except MilvusException as e:
            log.critical("Failed to connect to Milvus", error=str(e), exc_info=True)
            _milvus_connected = False
            raise ConnectionError(f"Milvus connection failed: {e}") from e

    if not _milvus_collection:
        try:
            if not utility.has_collection(MILVUS_COLLECTION_NAME, using=alias):
                log.warning(f"Milvus collection '{MILVUS_COLLECTION_NAME}' not found. Creating...")
                _create_milvus_collection(alias)
            else:
                log.info(f"Milvus collection '{MILVUS_COLLECTION_NAME}' exists.")

            _milvus_collection = Collection(name=MILVUS_COLLECTION_NAME, using=alias)
            log.info(f"Loading collection '{MILVUS_COLLECTION_NAME}' into memory...")
            _milvus_collection.load()
            log.info(f"Collection '{MILVUS_COLLECTION_NAME}' loaded.")

        except MilvusException as e:
            log.error(f"Failed to get or load Milvus collection '{MILVUS_COLLECTION_NAME}'", error=str(e), exc_info=True)
            _milvus_collection = None
            raise RuntimeError(f"Milvus collection error: {e}") from e

    return _milvus_collection

def _create_milvus_collection(alias: str):
    log.info(f"Defining schema for collection '{MILVUS_COLLECTION_NAME}'")
    fields = [
        FieldSchema(name=MILVUS_PK_FIELD, dtype=DataType.INT64, is_primary=True, auto_id=True),
        FieldSchema(name=MILVUS_VECTOR_FIELD, dtype=DataType.FLOAT_VECTOR, dim=MILVUS_EMBEDDING_DIM),
        FieldSchema(name=MILVUS_CONTENT_FIELD, dtype=DataType.VARCHAR, max_length=settings.SPLITTER_CHUNK_SIZE * 4),
        FieldSchema(name=MILVUS_COMPANY_ID_FIELD, dtype=DataType.VARCHAR, max_length=64),
        FieldSchema(name=MILVUS_DOCUMENT_ID_FIELD, dtype=DataType.VARCHAR, max_length=64),
        FieldSchema(name=MILVUS_FILENAME_FIELD, dtype=DataType.VARCHAR, max_length=512),
    ]
    schema = CollectionSchema(fields, description="Document Chunks for Atenex RAG")

    log.info(f"Creating collection '{MILVUS_COLLECTION_NAME}'...")
    try:
        collection = Collection(name=MILVUS_COLLECTION_NAME, schema=schema, using=alias, consistency_level="Strong")
        log.info(f"Collection '{MILVUS_COLLECTION_NAME}' created. Creating index...")

        index_params = settings.MILVUS_INDEX_PARAMS
        log.info("Creating index for vector field", field_name=MILVUS_VECTOR_FIELD, index_params=index_params)
        collection.create_index(field_name=MILVUS_VECTOR_FIELD, index_params=index_params)
        log.info("Index created successfully.")

        log.info("Creating scalar index for company_id field...")
        collection.create_index(field_name=MILVUS_COMPANY_ID_FIELD, index_name="company_id_idx")
        log.info("Creating scalar index for document_id field...")
        collection.create_index(field_name=MILVUS_DOCUMENT_ID_FIELD, index_name="document_id_idx")
        log.info("Scalar indexes created.")

    except MilvusException as e:
        log.error("Failed to create Milvus collection or index", collection_name=MILVUS_COLLECTION_NAME, error=str(e), exc_info=True)
        raise RuntimeError(f"Milvus collection/index creation failed: {e}") from e

def delete_milvus_chunks(company_id: str, document_id: str) -> int:
    del_log = log.bind(company_id=company_id, document_id=document_id)
    try:
        collection = _ensure_milvus_connection_and_collection()
        expr = f'{MILVUS_COMPANY_ID_FIELD} == "{company_id}" and {MILVUS_DOCUMENT_ID_FIELD} == "{document_id}"'
        del_log.info("Attempting to delete existing chunks from Milvus", filter_expr=expr)
        delete_result = collection.delete(expr=expr)
        deleted_count = delete_result.delete_count
        del_log.info("Milvus deletion successful", deleted_count=deleted_count)
        return deleted_count
    except MilvusException as e:
        del_log.error("Failed to delete chunks from Milvus", error=str(e), exc_info=True)
        raise RuntimeError(f"Milvus deletion failed: {e}") from e
    except Exception as e:
        del_log.exception("Unexpected error during Milvus chunk deletion", error=str(e))
        raise RuntimeError(f"Unexpected Milvus deletion error: {e}") from e


# --------------- 4. EMBEDDING MODEL INITIALIZATION REMOVED FROM GLOBAL SCOPE ---------------
# LLM_FLAG: REMOVED - Global embedding_model = TextEmbedding(...) block was here.

# --------------- 5. MAIN INGEST FUNCTION (MODIFIED) -----------------

# LLM_FLAG: MODIFIED - Added embedding_model argument to function signature
def ingest_document_pipeline(
    file_path: pathlib.Path,
    company_id: str,
    document_id: str,
    content_type: str,
    embedding_model: TextEmbedding, # Argument added
    delete_existing: bool = True
) -> int:
    """
    Processes a single document: extracts text, chunks, embeds, and inserts into Milvus.
    Args:
        file_path: Path to the downloaded document file.
        company_id: The company ID for multi-tenancy.
        document_id: The unique ID for the document.
        content_type: The MIME type of the file.
        embedding_model: The initialized FastEmbed TextEmbedding instance.
        delete_existing: If True, delete existing chunks for this company/document before inserting.
    Returns:
        The number of chunks successfully inserted into Milvus.
    Raises:
        ValueError: If the content type is unsupported, extraction fails, or model not provided.
        ConnectionError: If connection to Milvus fails.
        RuntimeError: For other Milvus or Embedding errors.
    """
    ingest_log = log.bind(
        company_id=company_id,
        document_id=document_id,
        filename=file_path.name,
        content_type=content_type
    )
    ingest_log.info("Starting ingestion pipeline for document")

    # --- 0. Check if embedding model was provided ---
    # LLM_FLAG: MODIFIED - Check the passed embedding_model argument
    if not embedding_model:
        ingest_log.error("Embedding model was not provided to the pipeline function.")
        raise ValueError("Embedding model instance is required for the pipeline.")

    # --- 1. Select Extractor ---
    extractor = EXTRACTORS.get(content_type)
    if not extractor:
        ingest_log.error("Unsupported content type for extraction.")
        raise ValueError(f"Unsupported content type: {content_type}")

    # --- 2. Extract Text ---
    ingest_log.debug("Extracting text content...")
    try:
        text_content = extractor(file_path)
        if not text_content or text_content.isspace():
            ingest_log.warning("No text content extracted from the document. Skipping.")
            return 0
        ingest_log.info(f"Text extracted successfully, length: {len(text_content)} chars.")
    except ValueError as ve:
        ingest_log.error("Text extraction failed.", error=str(ve))
        raise

    # --- 3. Chunk Text ---
    ingest_log.debug("Chunking extracted text...")
    chunks = chunk_text(
        text_content,
        chunk_size=settings.SPLITTER_CHUNK_SIZE,
        chunk_overlap=settings.SPLITTER_CHUNK_OVERLAP
    )
    if not chunks:
        ingest_log.warning("Text content resulted in zero chunks. Skipping.")
        return 0
    ingest_log.info(f"Text chunked into {len(chunks)} chunks.")

    # --- 4. Embed Chunks ---
    ingest_log.debug(f"Generating embeddings for {len(chunks)} chunks...")
    try:
        # LLM_FLAG: MODIFIED - Use the passed embedding_model argument
        embeddings = list(embedding_model.embed(chunks))
        ingest_log.info(f"Embeddings generated successfully for {len(embeddings)} chunks.")
        if len(embeddings) != len(chunks):
             ingest_log.warning("Mismatch between number of chunks and generated embeddings.", num_chunks=len(chunks), num_embeddings=len(embeddings))
             min_len = min(len(chunks), len(embeddings))
             chunks = chunks[:min_len]
             embeddings = embeddings[:min_len]

    except Exception as e:
        ingest_log.error("Failed to generate embeddings", error=str(e), exc_info=True)
        raise RuntimeError(f"Embedding generation failed: {e}") from e

    # --- 5. Prepare Data for Milvus ---
    data_to_insert = [
        embeddings,
        chunks,
        [company_id] * len(chunks),
        [document_id] * len(chunks),
        [file_path.name] * len(chunks),
    ]

    # --- 6. Delete Existing Chunks (Optional) ---
    if delete_existing:
        try:
            deleted_count = delete_milvus_chunks(company_id, document_id)
            ingest_log.info(f"Deleted {deleted_count} existing chunks before insertion.")
        except Exception as del_err:
            ingest_log.error("Failed to delete existing chunks, proceeding with insert anyway.", error=str(del_err))

    # --- 7. Insert into Milvus ---
    ingest_log.debug(f"Inserting {len(chunks)} chunks into Milvus collection '{MILVUS_COLLECTION_NAME}'...")
    try:
        collection = _ensure_milvus_connection_and_collection()
        mutation_result = collection.insert(data_to_insert)
        inserted_count = mutation_result.insert_count

        if inserted_count == len(chunks):
            ingest_log.info(f"Successfully inserted {inserted_count} chunks into Milvus.")
            log.debug("Flushing Milvus collection...")
            collection.flush()
            log.info("Milvus collection flushed.")
        else:
             ingest_log.warning(f"Milvus insert result count ({inserted_count}) differs from chunks sent ({len(chunks)}).")

        return inserted_count

    except MilvusException as e:
        ingest_log.error("Failed to insert data into Milvus", error=str(e), exc_info=True)
        raise RuntimeError(f"Milvus insertion failed: {e}") from e
    except Exception as e:
        ingest_log.exception("Unexpected error during Milvus insertion", error=str(e))
        raise RuntimeError(f"Unexpected Milvus insertion error: {e}") from e